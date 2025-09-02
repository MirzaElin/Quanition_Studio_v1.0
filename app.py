

from __future__ import annotations
import csv, json, math, os, sys, zipfile
from dataclasses import dataclass, field
from typing import Dict, List, Tuple, Optional, Any, Iterable, Set, DefaultDict
from collections import defaultdict, Counter

from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtCharts import QChart, QChartView, QLineSeries, QValueAxis

APP_TITLE = "Quanition Studio"
APP_VERSION = "v1.0"
BRAND = "Copyright © 2025 Mirza Niaz Zaman Elin. All Rights Reserved."


                                                   


@dataclass
class Dataset:
    name: str
    headers: List[str]
    rows: List[List[str]]

    @staticmethod
    def from_csv(path: str) -> "Dataset":
        with open(path, newline='', encoding='utf-8-sig') as f:
            reader = csv.reader(f)
            data = list(reader)
        if not data: raise ValueError("Empty CSV")
        headers = [str(x) for x in (data[0] or [])]
        rows = [[str(x) for x in r] for r in data[1:]]
        return Dataset(name=os.path.basename(path), headers=headers, rows=rows)

    @staticmethod
    def from_xlsx(path: str) -> "Dataset":
        import xml.etree.ElementTree as ET
        def _col_to_idx(cell_ref: str) -> int:
            letters = []
            for ch in cell_ref:
                if ch.isalpha(): letters.append(ch.upper())
                else: break
            idx = 0
            for ch in letters: idx = idx*26 + (ord(ch)-64)
            return max(0, idx-1)
        with zipfile.ZipFile(path, 'r') as z:
            shared = []
            if 'xl/sharedStrings.xml' in z.namelist():
                root = ET.fromstring(z.read('xl/sharedStrings.xml'))
                ns = {'a':'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
                for si in root.findall('.//a:si', ns):
                    texts = [t.text or '' for t in si.findall('.//a:t', ns)]
                    shared.append(''.join(texts))
            if 'xl/worksheets/sheet1.xml' in z.namelist():
                sheet_path='xl/worksheets/sheet1.xml'
            else:
                cands=[n for n in z.namelist() if n.startswith('xl/worksheets/') and n.endswith('.xml')]
                cands.sort()
                if not cands: raise ValueError("No worksheets found in .xlsx")
                sheet_path=cands[0]
            sheet_root=ET.fromstring(z.read(sheet_path)); ns={'a':'http://schemas.openxmlformats.org/spreadsheetml/2006/main'}
            table=[]; max_cols=0
            for row in sheet_root.findall('.//a:sheetData/a:row', ns):
                cells={}
                for c in row.findall('a:c', ns):
                    ref=c.get('r','A1'); j=_col_to_idx(ref); t=c.get('t'); v=c.find('a:v', ns)
                    val=''
                    if t=='s':
                        if v is not None and v.text is not None:
                            try: val=shared[int(v.text)]
                            except: val=''
                    elif t=='inlineStr':
                        tnode=c.find('.//a:t', ns); val=(tnode.text or '') if tnode is not None else ''
                    elif t=='b':
                        val='TRUE' if (v is not None and v.text=='1') else 'FALSE'
                    else:
                        if v is not None and v.text is not None: val=v.text
                        else:
                            tnode=c.find('.//a:t', ns); val=(tnode.text or '') if tnode is not None else ''
                    cells[j]=str(val); max_cols=max(max_cols, j+1)
                dense=[cells.get(j,'') for j in range(max_cols)]
                table.append(dense)
            if not table: raise ValueError("Empty worksheet")
            headers=[str(x) for x in table[0]]
            if all(h.strip()=='' for h in headers):
                headers=[f"col_{i+1}" for i in range(len(headers))]; rows=[[str(x) for x in r] for r in table]
            else:
                rows=[[str(x) for x in r] for r in table[1:]]
            return Dataset(name=os.path.basename(path), headers=headers, rows=rows)

    def to_table_model(self) -> QtGui.QStandardItemModel:
        model = QtGui.QStandardItemModel()
        model.setHorizontalHeaderLabels(self.headers)
        for r in self.rows:
            model.appendRow([QtGui.QStandardItem(str(x)) for x in r])
        return model

    def column(self, name: str) -> List[str]:
        try:
            j = self.headers.index(name)
            return [row[j] if j < len(row) else '' for row in self.rows]
        except ValueError:
            return []


                                                              


@dataclass
class Spec:
    framework: str
    config: Dict[str, Any] = field(default_factory=dict)

@dataclass
class RunResult:
    ok: bool
    summary: str
    figures: List[Tuple[str, List[Tuple[float, float]]]] = field(default_factory=list)
    tables: Dict[str, List[List[str]]] = field(default_factory=dict)
    details: Dict[str, Any] = field(default_factory=dict)


                                                                 


class HilbertModule:
    @staticmethod
    def run(cfg: Dict[str, Any]) -> RunResult:
        try:
            pA=float(cfg.get('pA',0.5)); pB_A=float(cfg.get('pB_given_A',0.5)); pA_B=float(cfg.get('pA_given_B',0.5))
            Ay=[1+0j,0+0j]; ampAy=math.sqrt(max(0.0,pA)); ampAn=math.sqrt(max(0.0,1-pA))
            theta=math.acos(max(0.0,min(1.0,math.sqrt(max(0.0,pB_A))))); By=[math.cos(theta)+0j, math.sin(theta)+0j]
            import cmath
            def pA_given_B(phi):
                psi=[ampAy, cmath.exp(1j*phi)*ampAn]
                amp_By=(By[0].conjugate()*psi[0] + By[1].conjugate()*psi[1])
                pb=abs(amp_By)**2
                if pb<=1e-12: return 0.0
                psi_By=[By[0]*amp_By/pb**0.5, By[1]*amp_By/pb**0.5]
                amp_Ay=(Ay[0].conjugate()*psi_By[0] + Ay[1].conjugate()*psi_By[1])
                return min(1.0,max(0.0,abs(amp_Ay)**2))
            best_phi, best_err=0.0, 1e9
            for k in range(720):
                phi=2*math.pi*k/720.0; val=pA_given_B(phi); err=abs(val-pA_B)
                if err<best_err: best_err, best_phi=err, phi
            est=pA_given_B(best_phi)
            series=[(float(k), pA_given_B(math.radians(k))) for k in range(0,361,6)]
            tbl=[["k_deg","pA|B(phi)"]]+[[f"{x:.1f}", f"{y:.4f}"] for x,y in series[::5]]
            return RunResult(True, f"Hilbert fit: φ≈{best_phi:.3f}, p(A|B)≈{est:.3f} (target {pA_B:.3f}); θ≈{theta:.3f}.",
                             figures=[("p(A|B) vs φ (deg)", series)], tables={"phi_sweep_sample": tbl},
                             details={"phi":best_phi,"theta":theta,"pA":pA,"pB|A":pB_A,"pA|B":pA_B})
        except Exception as e:
            return RunResult(False, f"Hilbert error: {e}")

class QDTModule:
    @staticmethod
    def softmax(us, tau=1.0):
        if tau<=0: m=max(us); z=[1.0 if u==m else 0.0 for u in us]; s=sum(z); return [zi/s for zi in z]
        exps=[math.exp(u/tau) for u in us]; Z=sum(exps); return [e/Z for e in exps]
    @staticmethod
    def run(cfg: Dict[str, Any]) -> RunResult:
        try:
            pros = cfg.get("prospects", [])
            tau=float(cfg.get("tau",1.0)); quarter=float(cfg.get("quarter",0.25))
            us=[float(p.get("utility",0.0)) for p in pros]; names=[str(p.get("name","")) for p in pros]
            uf=QDTModule.softmax(us,tau)
            freqs=[p.get("freq",None) for p in pros]
            if not pros: return RunResult(False, "QDT: no prospects provided.")
            if all(f is None for f in freqs):
                q=[(quarter if i%2==0 else -quarter) for i in range(len(pros))]; s=sum(q)
                if len(q)>=2: q[0]-=s
                ps=[max(0.0,min(1.0,uf[i]+q[i])) for i in range(len(q))]; Z=sum(ps); ps=[pi/Z for pi in ps]
            else:
                target=[float(f) if f is not None else uf[i] for i,f in enumerate(freqs)]
                q=[target[i]-uf[i] for i in range(len(target))]
                def proj(q): s=sum(q); q=[qi - s/len(q) for qi in q]; return [max(-1.0,min(1.0,qi)) for qi in q]
                for _ in range(100):
                    q=proj(q); avg=sum(abs(qi) for qi in q)/len(q) if q else 0.0
                    if avg==0: break
                    scale=quarter/avg if avg>0 else 0.0; q=[qi*0.5 + qi*0.5*scale for qi in q]
                ps=[max(0.0,min(1.0,uf[i]+q[i])) for i in range(len(q))]; Z=sum(ps); ps=[pi/Z for pi in ps]
            table=[["Prospect","Utility","Utility Factor","Final P"]]+[[names[i], f"{us[i]:.4f}", f"{uf[i]:.4f}", f"{ps[i]:.4f}"] for i in range(len(us))]
            series=[(float(i),ps[i]) for i in range(len(ps))]
            return RunResult(True, "QDT: utility + attraction factors computed.", figures=[("QDT P(choice) by prospect index", series)], tables={"qdt":table}, details={"final_probabilities":ps,"utilities":us,"names":names})
        except Exception as e:
            return RunResult(False, f"QDT error: {e}")

class CbDModule:
    @staticmethod
    def run(cfg: Dict[str, Any]) -> RunResult:
        try:
            E11=float(cfg.get('E11')); E21=float(cfg.get('E21')); E22=float(cfg.get('E22')); E12=float(cfg.get('E12'))
            mA1=float(cfg.get('mA1')); mA2=float(cfg.get('mA2')); mB1=float(cfg.get('mB1')); mB2=float(cfg.get('mB2'))
            Es=[E11,E21,E22,E12]
            signs=[[1,1,1,-1],[1,1,-1,1],[1,-1,1,1],[-1,1,1,1],[-1,-1,-1,1],[-1,-1,1,-1],[-1,1,-1,-1],[1,-1,-1,-1]]
            Sodd=max(sum(s[i]*Es[i] for i in range(4)) for s in signs); ICC=abs(mA1-mA2)+abs(mB1-mB2)
            contextual=Sodd>2+ICC
            table=[["Metric","Value"],["S_odd",f"{Sodd:.4f}"],["ICC",f"{ICC:.4f}"],["Threshold",f"{2+ICC:.4f}"],["Contextual?",str(contextual)]]
            return RunResult(True, f"CbD: S_odd={Sodd:.3f}, ICC={ICC:.3f}. Contextual? {contextual}.", figures=[("S_odd vs threshold", [(0.0,Sodd),(1.0,2+ICC)])], tables={"CbD":table}, details={"S_odd":Sodd,"ICC":ICC,"contextual":contextual})
        except Exception as e:
            return RunResult(False, f"CbD error: {e}")

class QGameModule:
    
    
    
    
    
    @staticmethod
    def _eisert_payoffs(R,S,T,P,gamma,a,b):
        
        I = [[1+0j,0+0j],[0+0j,1+0j]]
        X = [[0+0j,1+0j],[1+0j,0+0j]]
        Z = [[1+0j,0+0j],[0+0j,-1+0j]]
        def kron2(A,B):
            out = [[0+0j for _ in range(4)] for __ in range(4)]
            for i in range(2):
                for j in range(2):
                    for k in range(2):
                        for l in range(2):
                            out[2*i+k][2*j+l] = A[i][j]*B[k][l]
            return out
        def mat_vec(A,v):
            n=len(A); m=len(A[0]); out=[0+0j for _ in range(n)]
            for i in range(n):
                s=0+0j
                for j in range(m):
                    s+=A[i][j]*v[j]
                out[i]=s
            return out
        def dagger(A):
            n=len(A); m=len(A[0]); out=[[0+0j for _ in range(n)] for __ in range(m)]
            for i in range(n):
                for j in range(m):
                    out[j][i]=A[i][j].conjugate()
            return out
        
        epp = complex(math.cos(gamma/2), math.sin(gamma/2))
        emm = complex(math.cos(gamma/2),-math.sin(gamma/2))
        Jg = [[epp,0,0,0],[0,emm,0,0],[0,0,emm,0],[0,0,0,epp]]
        Jd = dagger(Jg)
        
        def strategy(tag):
            if tag=='C': return I
            if tag=='D': return [[1j*X[0][0],1j*X[0][1]],[1j*X[1][0],1j*X[1][1]]]  
            if tag=='Q': return [[1j*Z[0][0],1j*Z[0][1]],[1j*Z[1][0],1j*Z[1][1]]]  
            return I
        UA = strategy(a); UB = strategy(b)
        U = kron2(UA,UB)
        psi0 = [1+0j,0+0j,0+0j,0+0j]
        v = mat_vec(Jg, psi0)
        v = mat_vec(U, v)
        v = mat_vec(Jd, v)
        p00 = (v[0].real**2 + v[0].imag**2)
        p01 = (v[1].real**2 + v[1].imag**2)
        p10 = (v[2].real**2 + v[2].imag**2)
        p11 = (v[3].real**2 + v[3].imag**2)
        
        EA = p00*R + p01*S + p10*T + p11*P
        EB = p00*R + p01*T + p10*S + p11*P
        return EA, EB, (p00,p01,p10,p11)

    @staticmethod
    def run(cfg: Dict[str, Any]) -> RunResult:
        try:
            R=float(cfg.get('R',3)); S=float(cfg.get('S',0)); T=float(cfg.get('T',5)); P=float(cfg.get('P',1))
            gamma=float(cfg.get('gamma',0.0)); grid=int(cfg.get('grid',35))  
            strats=['C','D','Q']
            pay={}
            for a in strats:
                for b in strats:
                    EA,EB,_=QGameModule._eisert_payoffs(R,S,T,P,gamma,a,b)
                    pay[(a,b)]=(EA,EB)
            
            best_pair=None; best_val=-1e18; best_record=[]
            for a in strats:
                for b in strats:
                    EA,EB=pay[(a,b)]
                    s=EA+EB
                    if s>best_val+1e-12:
                        best_val=s; best_pair=(a,b); best_record=[(a,b,EA,EB)]
                    elif abs(s-best_val)<=1e-12:
                        best_record.append((a,b,EA,EB))
            
            table=[["strategy_A","strategy_B","EA","EB","EA+EB"]]
            for a,b,EA,EB in best_record:
                table.append([a,b,f"{EA:.4f}",f"{EB:.4f}",f"{(EA+EB):.4f}"])
            spectrum=[(float(i), pay[p][0]+pay[p][1]) for i,p in enumerate(pay)]
            
            return RunResult(
                True,
                f"QGame: max(EA+EB) = {best_val:.4f} at {best_pair}.",
                figures=[("EA+EB over strategy pairs (index order)", spectrum)],
                tables={"best_profiles": table},
                details={"best": best_record, "gamma": gamma, "payoffs": pay}
            )
        except Exception as e:
            return RunResult(False, f"QGame error: {e}")

class QFinanceModule:
    @staticmethod
    def binomial(S0,K,r,sigma,T,steps=100):
        dt=T/steps; u=math.exp(sigma*math.sqrt(dt)); d=1.0/u; p=(math.exp(r*dt)-d)/(u-d); p=max(0.0,min(1.0,p))
        disc=math.exp(-r*T); prices=[S0*(u**j)*(d**(steps-j)) for j in range(steps+1)]; pay=[max(0.0,s-K) for s in prices]
        for n in range(steps,0,-1):
            pay=[disc**(1.0/steps)*(p*pay[j+1]+(1-p)*pay[j]) for j in range(n)]
        return pay[0]
    @staticmethod
    def run(cfg: Dict[str, Any]) -> RunResult:
        try:
            S0=float(cfg.get('S0',100)); K=float(cfg.get('K',100)); r=float(cfg.get('r',0.01)); sigma=float(cfg.get('sigma',0.2)); T=float(cfg.get('T',1.0)); steps=int(cfg.get('steps',150))
            b=QFinanceModule.binomial(S0,K,r,sigma,T,steps)
            return RunResult(True, f"QFinance: Binomial call price ≈ {b:.4f}.", figures=[("Option price (binomial)", [(0.0,b),(1.0,b)])], tables={"prices":[["Method","Price"],["Binomial",f"{b:.4f}"]]}, details={"binomial":b})
        except Exception as e:
            return RunResult(False, f"QFinance error: {e}")

class QMoneyModule:
    @staticmethod
    def simulate(cfg: Dict[str, Any]):
        T=int(cfg.get('T',40)); c0=float(cfg.get('c0',1.0)); c1=float(cfg.get('c1',0.05)); rD=float(cfg.get('rD',0.01)); rL=float(cfg.get('rL',0.03)); inv0=float(cfg.get('inv0',0.5))
        cash_share=max(0.0,min(1.0,float(cfg.get('cash_share',0.5)))); corr=max(0.0,min(1.0,float(cfg.get('corr',0.2))))
        H=float(cfg.get('H0',100.0)); D=float(cfg.get('D0',50.0)); L=float(cfg.get('L0',20.0))
        out={k:[] for k in ["H","D","L","C","Y","cash","account"]}
        for t in range(T):
            C=c0+c1*H; Y=C; Hn=H+Y-C+rD*D-rL*L + corr*(D-L)*0.01; Dn=D+Y-C; Ln=max(0.0, L+inv0-(Y-C))
            cash=cash_share*Dn; account=(1-cash_share)*Dn
            H,D,L=Hn,Dn,Ln; out["H"].append(H); out["D"].append(D); out["L"].append(L); out["C"].append(C); out["Y"].append(Y); out["cash"].append(cash); out["account"].append(account)
        return out
    @staticmethod
    def run(cfg: Dict[str, Any]) -> RunResult:
        try:
            traj=QMoneyModule.simulate(cfg); series=[(float(i), traj["H"][i]) for i in range(len(traj["H"]))]
            table=[["t","H","D","L","C","Y","cash","account"]] + [[str(i), *(f"{traj[k][i]:.3f}" for k in ["H","D","L","C","Y","cash","account"])] for i in range(min(30,len(traj["H"])))]
            return RunResult(True, "Monetary sandbox simulated.", figures=[("Household wealth over time", series)], tables={"trajectory":table}, details=traj)
        except Exception as e:
            return RunResult(False, f"QMoney error: {e}")


                                                               


def _to_boolish(x: str) -> Optional[bool]:
    t = str(x).strip().lower()
    if t in ("yes","y","true","t","1"): return True
    if t in ("no","n","false","f","0"): return False
    return None

def _categorical_values(values: List[str], max_unique=10) -> Optional[List[str]]:
    uniq = list(dict.fromkeys([v.strip() for v in values if v.strip()!=""]))
    if len(uniq)<=max_unique: return uniq
    boolish = [_to_boolish(v) for v in values]
    if all(b is not None for b in boolish): return ["True","False"]
    return None

class AutoRunner:
    def __init__(self, ds: Dataset):
        self.ds = ds

    def run_hilbert_all(self, limit_pairs: int = 200) -> RunResult:
        headers = self.ds.headers
        cols = {h: self.ds.column(h) for h in headers}
        cats = {h: _categorical_values(cols[h]) for h in headers}
        cat_headers = [h for h in headers if cats.get(h)]
        results_rows = [["A_col","A_yes","B_col","B_yes","N","nA","nB","nAB","pA","pB|A","pA|B","phi","theta"]]
        count=0
        figures=[]
        examples_captured=0
        for i in range(len(cat_headers)):
            for j in range(i+1, len(cat_headers)):
                A, B = cat_headers[i], cat_headers[j]
                A_vals = [v for v in cols[A] if v.strip()!=""]; B_vals = [v for v in cols[B] if v.strip()!=""]
                if not A_vals or not B_vals: continue
                A_yes = Counter(A_vals).most_common(1)[0][0]; B_yes = Counter(B_vals).most_common(1)[0][0]
                N = 0; nA = 0; nB = 0; nAB = 0
                for a,b in zip(cols[A], cols[B]):
                    if a.strip()=="" or b.strip()=="": continue
                    N += 1
                    aY = (a==A_yes); bY = (b==B_yes)
                    if aY: nA += 1
                    if bY: nB += 1
                    if aY and bY: nAB += 1
                if N==0 or nA==0 or nB==0: continue
                pA = nA/N; pB_A = nAB/nA; pA_B = nAB/nB
                res = HilbertModule.run({"pA":pA, "pB_given_A":pB_A, "pA_given_B":pA_B})
                if res.ok:
                    phi = res.details.get("phi", 0.0); theta = res.details.get("theta", 0.0)
                    results_rows.append([A, A_yes, B, B_yes, f"{N}", f"{nA}", f"{nB}", f"{nAB}", f"{pA:.4f}", f"{pB_A:.4f}", f"{pA_B:.4f}", f"{phi:.4f}", f"{theta:.4f}"])
                    if examples_captured < 2 and res.figures:
                        title = f"Hilbert φ-sweep: {A}→{B}"
                        figures.append((title, res.figures[0][1]))
                        examples_captured += 1
                    count+=1
                    if count>=limit_pairs: break
            if count>=limit_pairs: break
        if len(results_rows)==1:
            return RunResult(False, "Hilbert Auto: no suitable categorical column pairs found.")
        return RunResult(True, f"Hilbert Auto: analyzed {len(results_rows)-1} column-pairs (one-vs-rest).", tables={"hilbert_pairs": results_rows}, figures=figures)

    def run_qdt_all(self) -> RunResult:
        hdrs = set(self.ds.headers)
        if not {"prospect","utility"}.issubset(hdrs):
            return RunResult(False, "QDT Auto: dataset needs 'prospect' and 'utility' columns (optional: 'freq','experiment','tau','quarter').")
        exp_col = "experiment" if "experiment" in hdrs else None
        rows = self.ds.rows; h_index = {h:i for i,h in enumerate(self.ds.headers)}
        buckets: DefaultDict[str, List[Dict[str,Any]]] = defaultdict(list)
        for r in rows:
            name = r[h_index["prospect"]] if h_index["prospect"]<len(r) else ""
            util = r[h_index["utility"]] if h_index["utility"]<len(r) else ""
            if name.strip()=="" or util.strip()=="": continue
            freq = r[h_index["freq"]] if "freq" in hdrs and h_index["freq"]<len(r) else None
            expv = r[h_index[exp_col]] if exp_col and h_index[exp_col]<len(r) else "default"
            try:
                utilf = float(util); freqf = (float(freq) if freq not in (None,"") else None)
                buckets[expv].append({"name":name,"utility":utilf,"freq":freqf})
            except:
                continue
        if not buckets: return RunResult(False, "QDT Auto: no valid rows parsed.")
        master_table = [["experiment","prospect","utility","utility_factor","final_P"]]
        details = {}
        figures=[]
        for exp, plist in buckets.items():
            res = QDTModule.run({"prospects": plist})
            if not res.ok: continue
            details[exp]=res.details
            tbl = res.tables.get("qdt",[])
            for row in tbl[1:]:
                master_table.append([exp, row[0], row[1], row[2], row[3]])
            if res.figures:
                title = f"QDT — {exp} (P by index)"
                figures.append((title, res.figures[0][1]))
        return RunResult(True, f"QDT Auto: processed {len(buckets)} experiment group(s).", tables={"qdt_all": master_table}, details=details, figures=figures)

    def run_cbd_all(self) -> RunResult:
        hdrs = set(self.ds.headers)
        req = {"A1","A2","B1","B2"}
        if not req.issubset(hdrs):
            return RunResult(False, "CbD Auto: needs columns A1,A2,B1,B2 with ±1 values.")
        cols = {h: [float(x) if str(x).strip()!="" else 0.0 for x in self.ds.column(h)] for h in req}
        def mean(vs): 
            xs=[x for x in vs if x in (-1.0,1.0)]
            return sum(xs)/len(xs) if xs else 0.0
        def mean_prod(a,b):
            xs=[ai*bi for ai,bi in zip(cols[a], cols[b]) if ai in (-1.0,1.0) and bi in (-1.0,1.0)]
            return sum(xs)/len(xs) if xs else 0.0
        vals = {
            "E11": mean_prod("A1","B1"), "E12": mean_prod("A1","B2"),
            "E21": mean_prod("A2","B1"), "E22": mean_prod("A2","B2"),
            "mA1": mean(cols["A1"]), "mA2": mean(cols["A2"]),
            "mB1": mean(cols["B1"]), "mB2": mean(cols["B2"]),
        }
        res = CbDModule.run(vals)
        tbl = [["Param","Value"]] + [[k,f"{v:.4f}"] for k,v in vals.items()]
        out_tables = {"inputs": tbl}
        out_tables.update(res.tables)
        figs = res.figures if res.figures else []
        return RunResult(True if res.ok else False, f"CbD Auto: computed from A1,A2,B1,B2.", tables=out_tables, figures=figs, details={"inputs":vals})

    def run_qgame_all(self, max_rows: int = 50) -> RunResult:
        need = ["R","S","T","P"]
        miss = [k for k in need if k not in self.ds.headers]
        if miss:
            return RunResult(False, f"QGame Auto: missing columns {', '.join(miss)} (optional: gamma, grid).")
        idx = {h:i for i,h in enumerate(self.ds.headers)}
        table=[["row","R","S","T","P","gamma","grid","EA+EB_max","at_strategy"]]
        n=0
        figures=[]
        captured_example=False
        for r in self.ds.rows:
            try:
                R=float(r[idx["R"]]); S=float(r[idx["S"]]); T=float(r[idx["T"]]); P=float(r[idx["P"]])
                gamma=float(r[idx["gamma"]]) if "gamma" in idx and r[idx["gamma"]].strip()!="" else 0.6
                grid=int(float(r[idx["grid"]])) if "grid" in idx and r[idx["grid"]].strip()!="" else 35
                res=QGameModule.run({"R":R,"S":S,"T":T,"P":P,"gamma":gamma,"grid":grid})
                if res.ok and res.details.get("best"):
                    best = res.details["best"]
                    best_val = max(EA+EB for (_,_,EA,EB) in best)
                    at = ", ".join([f"({a},{b})" for (a,b,_,_) in best])
                    table.append([str(n), f"{R}", f"{S}", f"{T}", f"{P}", f"{gamma}", f"{grid}", f"{best_val:.4f}", at])
                    if (not captured_example) and res.figures:
                        figures.append((f"QGame EA+EB spectrum (γ={gamma:.2f})", res.figures[0][1]))
                        captured_example=True
                    n+=1
                    if n>=max_rows: break
            except Exception as e:
                continue
        if n==0: return RunResult(False, "QGame Auto: no valid rows parsed.")
        return RunResult(True, f"QGame Auto: processed {n} scenario row(s).", tables={"qgame_scenarios": table}, figures=figures)

    def run_qfinance_all(self, max_rows: int = 200) -> RunResult:
        need = ["S0","K","r","sigma","T"]
        miss = [k for k in need if k not in self.ds.headers]
        if miss:
            return RunResult(False, f"QFinance Auto: missing columns {', '.join(miss)} (optional: steps).")
        idx = {h:i for i,h in enumerate(self.ds.headers)}
        table=[["row","S0","K","r","sigma","T","steps","price"]]
        n=0
        figures=[]
        captured=False
        for r in self.ds.rows:
            try:
                S0=float(r[idx["S0"]]); K=float(r[idx["K"]]); rr=float(r[idx["r"]]); sig=float(r[idx["sigma"]]); TT=float(r[idx["T"]])
                steps=int(float(r[idx["steps"]])) if "steps" in idx and r[idx["steps"]].strip()!="" else 150
                res=QFinanceModule.run({"S0":S0,"K":K,"r":rr,"sigma":sig,"T":TT,"steps":steps})
                if res.ok:
                    table.append([str(n), f"{S0}", f"{K}", f"{rr}", f"{sig}", f"{TT}", f"{steps}", f"{res.details['binomial']:.6f}"])
                    if (not captured) and res.figures:
                        figures.append(("QFinance example", res.figures[0][1]))
                        captured=True
                    n+=1
                    if n>=max_rows: break
            except:
                continue
        if n==0: return RunResult(False, "QFinance Auto: no valid rows parsed.")
        return RunResult(True, f"QFinance Auto: priced {n} option row(s).", tables={"qfinance_batch": table}, figures=figures)

    def run_qmoney_all(self, max_rows: int = 20) -> RunResult:
        need = ["T","c0","c1","rD","rL","inv0","cash_share","corr","H0","D0","L0"]
        if not any(k in self.ds.headers for k in need):
            return RunResult(False, "QMoney Auto: expects columns T,c0,c1,rD,rL,inv0,cash_share,corr,H0,D0,L0.")
        idx = {h:i for i,h in enumerate(self.ds.headers)}
        table=[["row","T","c0","c1","rD","rL","inv0","cash_share","corr","H0","D0","L0","H_T"]]
        n=0; fig_series=None
        for r in self.ds.rows:
            try:
                cfg={}
                for k in need:
                    if k in idx and idx[k]<len(r) and str(r[idx[k]]).strip()!="":
                        cfg[k]=float(r[idx[k]])
                req_min=["T","c0","c1","H0","D0","L0"]
                if not all(k in cfg for k in req_min): continue
                res=QMoneyModule.run(cfg)
                if res.ok:
                    lastH = res.details["H"][-1] if "H" in res.details and res.details["H"] else float('nan')
                    rowout=[str(n)] + [str(int(cfg.get("T",40)))]+[f"{cfg.get('c0',1.0)}",f"{cfg.get('c1',0.05)}",f"{cfg.get('rD',0.01)}",f"{cfg.get('rL',0.03)}",f"{cfg.get('inv0',0.5)}",f"{cfg.get('cash_share',0.5)}",f"{cfg.get('corr',0.2)}",f"{cfg.get('H0',100.0)}",f"{cfg.get('D0',50.0)}",f"{cfg.get('L0',20.0)}", f"{lastH:.3f}"]
                    table.append(rowout); n+=1
                    if fig_series is None: fig_series=res.figures[0][1] if res.figures else None
                    if n>=max_rows: break
            except:
                continue
        figs=[]
        if fig_series: figs=[("Example trajectory (first scenario)", fig_series)]
        if n==0: return RunResult(False, "QMoney Auto: no valid rows parsed.")
        return RunResult(True, f"QMoney Auto: simulated {n} scenario row(s).", tables={"qmoney_batch": table}, figures=figs)


                                                               


class ReportBuilder:
    @staticmethod
    def to_html(path: str, spec: Spec, result: RunResult):
        def esc(x): return x.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
        html=["<html><head><meta charset='utf-8'><title>QuanEcon Report</title>",
              "<style>body{font-family:system-ui,Segoe UI,Arial;margin:24px;}table{border-collapse:collapse;margin:12px 0;}th,td{border:1px solid #ddd;padding:6px 8px;}th{background:#f5f5f5}h1{margin:0}h2{margin-top:24px}footer{margin-top:24px;color:#666}</style>",
              "</head><body>"]
        html.append(f"<h1>{esc(APP_TITLE)} {esc(APP_VERSION)}</h1><div><em>{esc(BRAND)}</em></div>")
        html.append(f"<h2>Framework: {esc(spec.framework)}</h2>")
        html.append("<h3>Configuration</h3><pre>"+esc(json.dumps(spec.config, indent=2))+"</pre>")
        html.append(f"<h3>Summary</h3><p>{esc(result.summary)}</p>")
        for name, rows in result.tables.items():
            html.append(f"<h3>Table — {esc(name)}</h3><table>")
            for r, row in enumerate(rows):
                cell_tag = 'th' if r == 0 else 'td'
                row_html = "".join([f"<{cell_tag}>{esc(str(c))}</{cell_tag}>" for c in row])
                html.append("<tr>" + row_html + "</tr>")
            html.append("</table>")
        for title, series in result.figures:
            html.append(f"<h3>Figure — {esc(title)}</h3><pre>index,x,y\n")
            for i,(x,y) in enumerate(series): html.append(esc(f"{i},{x},{y}")+"\n")
            html.append("</pre>")
        html.append(f"<footer>{esc(BRAND)}</footer></body></html>")
        with open(path,"w",encoding="utf-8") as f: f.write("".join(html))

    @staticmethod
    def to_docx(path: str, spec: Spec, result: RunResult):
        
        
        try:
            from docx import Document
            d = Document()
            d.add_heading(f"{APP_TITLE} {APP_VERSION}", level=0)
            d.add_paragraph(BRAND)
            d.add_heading("Framework", level=1)
            d.add_paragraph(spec.framework)
            d.add_heading("Configuration", level=1)
            d.add_paragraph(json.dumps(spec.config, indent=2))
            d.add_heading("Summary", level=1)
            d.add_paragraph(result.summary)
            for name, rows in result.tables.items():
                d.add_heading(f"Table — {name}", level=2)
                if not rows:
                    continue
                table = d.add_table(rows=1, cols=len(rows[0]))
                hdr_cells = table.rows[0].cells
                for j, h in enumerate(rows[0]):
                    hdr_cells[j].text = str(h)
                for row in rows[1:]:
                    row_cells = table.add_row().cells
                    for j, val in enumerate(row):
                        row_cells[j].text = str(val)
            for title, series in result.figures:
                d.add_heading(f"Figure — {title}", level=2)
                para = d.add_paragraph()
                para.add_run("index,x,y\n").bold = True
                for i, (x, y) in enumerate(series[:500]):
                    para.add_run(f"{i},{x},{y}\n")
            d.add_paragraph("\n")
            d.add_paragraph(BRAND)
            d.save(path)
            return
        except Exception:
            pass
        html_path = path if path.lower().endswith(".doc") else (path + ".doc")
        ReportBuilder.to_html(html_path, spec, result)




                                

EULA_TEXT='QuanEcon Studio EULA — Contact: niazmirza111@gmail.com'
PRIVACY_TEXT='No telemetry. Files stay local.'
REQUIREMENTS_TEXT='Python 3.9+, PySide6, optional python-docx.'
THIRD_PARTY_TEXT='PySide6 (LGPL), Python (PSF).'
README_TEXT='QuanEcon Studio v1.0 “Masterpiece” build.'
USER_MANUAL_TEXT='Use Data → import, Run, Results → Export.'
SAMPLES_README='Includes CSV demos for all modules.'
def _mk_samples_dir():
    import tempfile, os
    d = os.path.join(tempfile.gettempdir(), 'quanecon_samples'); os.makedirs(d, exist_ok=True); return d
def write_sample_files():
    d = _mk_samples_dir(); files={}
    files['hilbert_demo.csv'] = 'A,B\nYes,No\nYes,Yes\nNo,Yes\nYes,Yes\nNo,No\nYes,No\n'
    files['qdt_demo.csv'] = 'experiment,prospect,utility,freq\nsnacks,Apple,0.10,0.35\nsnacks,Chocolate,0.25,0.45\nsnacks,Nuts,0.15,0.20\ngadgets,Phone X,0.30,0.50\ngadgets,Phone Y,0.25,0.35\ngadgets,Feature Phone,0.05,0.15\n'
    files['cbd_demo.csv'] = 'A1,A2,B1,B2\n1,1,1,1\n1,-1,1,-1\n-1,1,-1,1\n-1,-1,-1,-1\n'
    files['qfinance_demo.csv'] = 'S0,K,r,sigma,T,steps\n100,100,0.01,0.20,1.0,150\n120,100,0.02,0.25,0.5,120\n'
    files['qmoney_demo.csv'] = 'T,c0,c1,rD,rL,inv0,cash_share,corr,H0,D0,L0\n40,1,0.05,0.01,0.03,0.5,0.5,0.2,100,50,20\n'
    files['qgame_demo.csv'] = 'R,S,T,P,gamma,grid\n3,0,5,1,0.6,35\n'
    import os
    for name, content in files.items():
        with open(os.path.join(d, name), 'w', encoding='utf-8') as f: f.write(content)
    with open(os.path.join(d, 'README_samples.txt'), 'w', encoding='utf-8') as f: f.write(SAMPLES_README)
    return d
class DocsDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent); self.setWindowTitle('QuanEcon Docs'); self.resize(760,640)
        v=QtWidgets.QVBoxLayout(self); tabs=QtWidgets.QTabWidget(); v.addWidget(tabs)
        def mk(txt):
            w=QtWidgets.QWidget(); l=QtWidgets.QVBoxLayout(w); t=QtWidgets.QPlainTextEdit(); t.setReadOnly(True); t.setPlainText(txt); l.addWidget(t); return w
        tabs.addTab(mk(README_TEXT),'Read Me'); tabs.addTab(mk(USER_MANUAL_TEXT),'User Manual'); tabs.addTab(mk(REQUIREMENTS_TEXT),'Requirements'); tabs.addTab(mk(THIRD_PARTY_TEXT),'Third Party'); tabs.addTab(mk(EULA_TEXT),'EULA'); tabs.addTab(mk(PRIVACY_TEXT),'Privacy')
        b=QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Close); b.rejected.connect(self.reject); v.addWidget(b)
LEARN_HILBERT='Hilbert: φ and θ control interference.'
LEARN_QDT='QDT: utilities + attraction.'
LEARN_CBD='CbD: S_odd vs (2+ICC).'
LEARN_QGAME='Eisert: {C,D,Q}×{C,D,Q}, entanglement γ.'
LEARN_QFIN='Binomial pricing (risk-neutral).'
LEARN_QMONEY='Toy monetary sandbox.'
class LearningCenterDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent); self.setWindowTitle('Learning Center — QuanEcon Studio'); self.resize(760,640)
        v=QtWidgets.QVBoxLayout(self); tabs=QtWidgets.QTabWidget(); v.addWidget(tabs)
        def mk(t):
            w=QtWidgets.QWidget(); l=QtWidgets.QVBoxLayout(w); p=QtWidgets.QPlainTextEdit(); p.setReadOnly(True); p.setPlainText(t); l.addWidget(p); return w
        tabs.addTab(mk(LEARN_HILBERT),'Hilbert'); tabs.addTab(mk(LEARN_QDT),'QDT'); tabs.addTab(mk(LEARN_CBD),'CbD'); tabs.addTab(mk(LEARN_QGAME),'QGame'); tabs.addTab(mk(LEARN_QFIN),'QFinance'); tabs.addTab(mk(LEARN_QMONEY),'QMoney')
        b=QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.Close); b.rejected.connect(self.reject); v.addWidget(b)
def open_docs_dialog(parent): dlg=DocsDialog(parent); dlg.exec()
def open_learning_dialog(parent): dlg=LearningCenterDialog(parent); dlg.exec()
def open_samples_menu(parent, loader):
    d=write_sample_files(); items=[('Hilbert Demo',os.path.join(d,'hilbert_demo.csv')),('QDT Demo',os.path.join(d,'qdt_demo.csv')),('CbD Demo',os.path.join(d,'cbd_demo.csv')),('QFinance Demo',os.path.join(d,'qfinance_demo.csv')),('QMoney Demo',os.path.join(d,'qmoney_demo.csv')),('QGame Demo',os.path.join(d,'qgame_demo.csv'))]
    labels=[x[0] for x in items]; choice,ok=QtWidgets.QInputDialog.getItem(parent,'Open Sample Dataset','Choose:',labels,0,False)
    if not ok: return
    sel=next((p for (lbl,p) in items if lbl==choice),None)
    if sel: loader(sel)

                                                                          


class ChartWidget(QtWidgets.QWidget):
    def __init__(self, title, seriesXY, parent=None):
        super().__init__(parent)
        layout = QtWidgets.QVBoxLayout(self)

        
        if not seriesXY or all((x is None or y is None) for x, y in seriesXY):
            lbl = QtWidgets.QLabel(f"{title}\n(No data to display)")
            lbl.setAlignment(QtCore.Qt.AlignCenter)
            layout.addWidget(lbl)
            return

        
        ls = QLineSeries()
        xmin = ymin = float("inf")
        xmax = ymax = float("-inf")
        for x, y in seriesXY:
            try:
                xf = float(x); yf = float(y)
            except Exception:
                continue
            ls.append(xf, yf)
            xmin = min(xmin, xf); xmax = max(xmax, xf)
            ymin = min(ymin, yf); ymax = max(ymax, yf)

        chart = QChart()
        chart.setTitle(title)
        chart.addSeries(ls)

        axX, axY = QValueAxis(), QValueAxis()
        axX.setTitleText("x")
        axY.setTitleText("y")

        def _pad(lo, hi):
            if not (math.isfinite(lo) and math.isfinite(hi)):
                return 0.0, 1.0
            if hi == lo:
                eps = 1.0 if lo == 0.0 else abs(lo) * 0.05
                return lo - eps, hi + eps
            span = hi - lo
            pad = span * 0.07
            return lo - pad, hi + pad

        xlo, xhi = _pad(xmin, xmax)
        ylo, yhi = _pad(ymin, ymax)

        axX.setRange(xlo, xhi)
        axY.setRange(ylo, yhi)

        chart.addAxis(axX, QtCore.Qt.AlignBottom)
        chart.addAxis(axY, QtCore.Qt.AlignLeft)
        ls.attachAxis(axX)
        ls.attachAxis(axY)

        view = QChartView(chart)
        view.setRenderHint(QtGui.QPainter.Antialiasing)
        layout.addWidget(view)

class TableWidget(QtWidgets.QTableWidget):
    def __init__(self, rows, parent=None):
        super().__init__(parent)
        if not rows: self.setRowCount(0); self.setColumnCount(0); return
        headers=rows[0]; self.setColumnCount(len(headers)); self.setHorizontalHeaderLabels(headers); self.setRowCount(len(rows)-1)
        for i,row in enumerate(rows[1:]):
            for j,val in enumerate(row): self.setItem(i,j,QtWidgets.QTableWidgetItem(str(val)))
        self.resizeColumnsToContents(); self.setMinimumHeight(min(400,30*(len(rows))))

class ResultsPanel(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__(parent); v=QtWidgets.QVBoxLayout(self); self.summary=QtWidgets.QLabel(""); self.summary.setWordWrap(True); v.addWidget(self.summary)
        self.scroll=QtWidgets.QScrollArea(); self.scroll.setWidgetResizable(True); self.container=QtWidgets.QWidget(); self.box=QtWidgets.QVBoxLayout(self.container)
        self.scroll.setWidget(self.container); v.addWidget(self.scroll)
        self.btn=QtWidgets.QPushButton("Export HTML Report…"); self.btn.clicked.connect(self.export_report); v.addWidget(self.btn)
        self.btnDocx=QtWidgets.QPushButton("Export DOCX Report…"); self.btnDocx.clicked.connect(self.export_docx); v.addWidget(self.btnDocx)
        self.spec=None; self.result=None
    def show_result(self, spec, result):
        self.spec,self.result=spec,result
        while self.box.count():
            it=self.box.takeAt(0); w=it.widget(); 
            if w: w.setParent(None)
        self.summary.setText(result.summary)
        for name,rows in result.tables.items():
            self.box.addWidget(QtWidgets.QLabel(f"<b>Table — {name}</b>")); self.box.addWidget(TableWidget(rows))
        for title,series in result.figures:
            self.box.addWidget(QtWidgets.QLabel(f"<b>Figure — {title}</b>")); self.box.addWidget(ChartWidget(title, series))
        self.box.addStretch(1)
    def export_report(self):
        if not (self.spec and self.result):
            QtWidgets.QMessageBox.information(self,"No result","Run a model first."); return
        path,_=QtWidgets.QFileDialog.getSaveFileName(self,"Save Report","quanecon_report.html","HTML (*.html)")
        if not path: return
        try:
            ReportBuilder.to_html(path, self.spec, self.result); QtWidgets.QMessageBox.information(self,"Saved",f"Report saved to:\n{path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,"Error",str(e))

    def export_docx(self):
        if not (self.spec and self.result):
            QtWidgets.QMessageBox.information(self,"No result","Run a model first."); return
        path,_=QtWidgets.QFileDialog.getSaveFileName(self,"Save DOCX Report","quanecon_report.docx","Word Document (*.docx);;Word 97-2003 (*.doc)")
        if not path: return
        try:
            ReportBuilder.to_docx(path, self.spec, self.result)
            QtWidgets.QMessageBox.information(self,"Saved",f"Report saved to:\n{path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,"Error",str(e))

class DataPanel(QtWidgets.QWidget):
    def __init__(self,parent=None):
        super().__init__(parent); v=QtWidgets.QVBoxLayout(self)
        self.loadBtn=QtWidgets.QPushButton("Import CSV/Excel…"); self.loadBtn.clicked.connect(self.load); v.addWidget(self.loadBtn)
        self.table=QtWidgets.QTableView(); v.addWidget(self.table); self.dataset=None
    def load(self):
        path,_=QtWidgets.QFileDialog.getOpenFileName(self,"Open Data","","Data Files (*.csv *.xlsx);;CSV (*.csv);;Excel (*.xlsx)")
        if not path: return
        try:
            if path.lower().endswith(".xlsx"): ds=Dataset.from_xlsx(path)
            else: ds=Dataset.from_csv(path)
            self.dataset=ds; self.table.setModel(ds.to_table_model())
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,"Error",str(e))

    def load_path(self, path: str):
        if not path: return
        try:
            if path.lower().endswith('.xlsx'):
                ds=Dataset.from_xlsx(path)
            else:
                ds=Dataset.from_csv(path)
            self.dataset=ds; self.table.setModel(ds.to_table_model())
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,'Error',str(e))

class SpecEditor(QtWidgets.QWidget):
    specChanged = QtCore.Signal(Spec)
    def __init__(self,parent=None):
        super().__init__(parent); f=QtWidgets.QFormLayout(self)
        self.box=QtWidgets.QComboBox(); self.box.addItems(["Hilbert","QDT","CbD","QGame","QFinance","QMoney"]); self.box.currentTextChanged.connect(self.template)
        self.text=QtWidgets.QPlainTextEdit(); self.text.setPlaceholderText("JSON configuration here…")
        btn=QtWidgets.QPushButton("Validate Spec"); btn.clicked.connect(self.emit)
        f.addRow("Framework", self.box); f.addRow(self.text); f.addRow(btn); self.template(self.box.currentText())
    def template(self,fw):
        if fw=="QDT":
            t = {
                "tau": 1.0,
                "quarter": 0.25,
                "experiments": [
                    { "label": "snacks",
                      "prospects": [
                        {"name":"Apple","utility":0.10,"freq":0.35},
                        {"name":"Chocolate","utility":0.25,"freq":0.45},
                        {"name":"Nuts","utility":0.15,"freq":0.20}
                      ]
                    },
                    { "label": "gadgets",
                      "prospects": [
                        {"name":"Phone X","utility":0.30,"freq":0.50},
                        {"name":"Phone Y","utility":0.25,"freq":0.35},
                        {"name":"Feature Phone","utility":0.05,"freq":0.15}
                      ]
                    }
                ]
            }
        else:
            t={"Hilbert":{"pA":0.55,"pB":0.52,"pA_given_B":0.60,"pB_given_A":0.58},
               "CbD":{"E11":0.7,"E21":0.6,"E22":0.6,"E12":0.7,"mA1":0.2,"mA2":0.1,"mB1":0.05,"mB2":0.02},
               "QGame":{"R":3,"S":0,"T":5,"P":1,"gamma":0.6,"grid":35},
               "QFinance":{"S0":100,"K":100,"r":0.01,"sigma":0.2,"T":1.0,"steps":150},
               "QMoney":{"T":40,"c0":1.0,"c1":0.05,"rD":0.01,"rL":0.03,"inv0":0.5,"cash_share":0.5,"corr":0.2,"H0":100,"D0":50,"L0":20}}
            t = t.get(fw,{})
        self.text.setPlainText(json.dumps(t, indent=2)); self.emit()
    def emit(self):
        try:
            cfg=json.loads(self.text.toPlainText() or "{}"); self.specChanged.emit(Spec(framework=self.box.currentText(), config=cfg))
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,"Invalid JSON",str(e))

class AutoPanel(QtWidgets.QWidget):
    autoFinished = QtCore.Signal(RunResult)
    def __init__(self, get_dataset_callable, parent=None):
        super().__init__(parent); self.get_dataset=get_dataset_callable
        v=QtWidgets.QVBoxLayout(self)
        v.addWidget(QtWidgets.QLabel("<b>Auto mode</b> — runs frameworks directly from your imported dataset."))
        self.btnRun = QtWidgets.QPushButton("Run All (Auto-detect)")
        self.btnRun.clicked.connect(self.run_all)
        v.addWidget(self.btnRun)
        self.log = QtWidgets.QTextEdit(); self.log.setReadOnly(True); self.log.setMinimumHeight(160); v.addWidget(self.log)
    def append(self, msg): self.log.append(msg)
    def run_all(self):
        ds = self.get_dataset()
        if not ds:
            QtWidgets.QMessageBox.information(self,"No data","Import a CSV/Excel first (Data tab)."); return
        self.append(f"Dataset: {ds.name} ({len(ds.rows)} rows, {len(ds.headers)} columns)")
        auto = AutoRunner(ds)
        tables={}
        figures=[]
        
        rH = auto.run_hilbert_all(limit_pairs=200); self.append(rH.summary)
        if rH.ok: tables.update(rH.tables); figures.extend(rH.figures)
        
        rQDT = auto.run_qdt_all(); self.append(rQDT.summary)
        if rQDT.ok: tables.update(rQDT.tables); figures.extend(rQDT.figures)
        
        rC = auto.run_cbd_all(); self.append(rC.summary)
        if rC.ok: tables.update(rC.tables); figures.extend(rC.figures)
        
        rG = auto.run_qgame_all(); self.append(rG.summary)
        if rG.ok: tables.update(rG.tables); figures.extend(rG.figures)
        
        rF = auto.run_qfinance_all(); self.append(rF.summary)
        if rF.ok: tables.update(rF.tables); figures.extend(rF.figures)
        
        rM = auto.run_qmoney_all(); self.append(rM.summary)
        if rM.ok: tables.update(rM.tables); figures.extend(rM.figures)
        all_ok = any([r.ok for r in [rH,rQDT,rC,rG,rF,rM]])
        total_tables = sum(len(t) for t in tables.values())
        res = RunResult(
            all_ok,
            f"Auto mode run complete: {len(tables)} table(s), {total_tables} row-block(s), {len(figures)} figure(s).",
            figures=figures, tables=tables,
            details={
                "hilbert": rH.details if rH.ok else {},
                "qdt":     rQDT.details if rQDT.ok else {},
                "cbd":     rC.details if rC.ok else {},
                "qgame":   rG.details if rG.ok else {},
                "qfinance":rF.details if rF.ok else {},
                "qmoney":  rM.details if rM.ok else {}
            }
        )
        self.autoFinished.emit(res)

class Runner(QtCore.QObject):
    finished = QtCore.Signal(RunResult)
    def __init__(self,spec:Spec): super().__init__(); self.spec=spec
    def run(self):
        fw=self.spec.framework; cfg=self.spec.config
        if fw=="QDT":
            
            tau=float(cfg.get("tau",1.0)); quarter=float(cfg.get("quarter",0.25))
            figures=[]; details={}; master=[["experiment","prospect","utility","utility_factor","final_P"]]
            
            if "experiments" in cfg and isinstance(cfg["experiments"], list):
                ok_any=False
                for ex in cfg["experiments"]:
                    label = str(ex.get("label","exp"))
                    pros = ex.get("prospects", [])
                    res = QDTModule.run({"prospects": pros, "tau": tau, "quarter": quarter})
                    if res.ok:
                        ok_any=True; details[label]=res.details
                        tbl = res.tables.get("qdt", [])
                        for row in tbl[1:]:
                            master.append([label, row[0], row[1], row[2], row[3]])
                        if res.figures:
                            figures.append((f"QDT — {label} (P by index)", res.figures[0][1]))
                if not ok_any:
                    self.finished.emit(RunResult(False, "QDT: no valid experiments.", figures=[], tables={})); return
                self.finished.emit(RunResult(True, f"QDT manual: processed {len(details)} experiment(s).", figures=figures, tables={"qdt_manual_all": master}, details=details)); return
            
            if "prospects" in cfg and isinstance(cfg["prospects"], list) and any("experiment" in p for p in cfg["prospects"]):
                buckets: DefaultDict[str, List[Dict[str,Any]]] = defaultdict(list)
                for p in cfg["prospects"]:
                    label=str(p.get("experiment","exp")); q={k:p[k] for k in p if k!="experiment"}
                    buckets[label].append(q)
                ok_any=False
                for label, pros in buckets.items():
                    res = QDTModule.run({"prospects": pros, "tau": tau, "quarter": quarter})
                    if res.ok:
                        ok_any=True; details[label]=res.details
                        tbl = res.tables.get("qdt", [])
                        for row in tbl[1:]:
                            master.append([label, row[0], row[1], row[2], row[3]])
                        if res.figures:
                            figures.append((f"QDT — {label} (P by index)", res.figures[0][1]))
                if not ok_any:
                    self.finished.emit(RunResult(False, "QDT: no valid prospects with 'experiment'.", figures=[], tables={})); return
                self.finished.emit(RunResult(True, f"QDT manual: processed {len(details)} experiment(s).", figures=figures, tables={"qdt_manual_all": master}, details=details)); return
            
            res = QDTModule.run(cfg); self.finished.emit(res); return
        
        if   fw=="Hilbert":  res=HilbertModule.run(cfg)
        elif fw=="CbD":      res=CbDModule.run(cfg)
        elif fw=="QGame":    res=QGameModule.run(cfg)
        elif fw=="QFinance": res=QFinanceModule.run(cfg)
        elif fw=="QMoney":   res=QMoneyModule.run(cfg)
        else:                res=RunResult(False, f"Unknown framework: {fw}")
        self.finished.emit(res)

class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__(); self.setWindowTitle(f"{APP_TITLE} {APP_VERSION} — {BRAND}"); self.resize(1280,860)
        self.dataset=None; self.spec=Spec(framework="Hilbert", config={}); self.result=None
        split=QtWidgets.QSplitter(); split.setOrientation(QtCore.Qt.Horizontal); self.setCentralWidget(split)
        left=QtWidgets.QWidget(); lv=QtWidgets.QVBoxLayout(left)
        self.btnHome=QtWidgets.QPushButton("Home"); self.btnData=QtWidgets.QPushButton("Data"); self.btnSpec=QtWidgets.QPushButton("Model Spec")
        self.btnAuto=QtWidgets.QPushButton("Auto"); self.btnRun=QtWidgets.QPushButton("Run Model"); self.btnResults=QtWidgets.QPushButton("Results")
        for b in [self.btnHome,self.btnData,self.btnSpec,self.btnAuto,self.btnRun,self.btnResults]: b.setMinimumHeight(40); lv.addWidget(b)
        lv.addStretch(1); split.addWidget(left)
        self.stack=QtWidgets.QStackedWidget(); split.addWidget(self.stack); split.setStretchFactor(1,1)
        self.home=QtWidgets.QWidget(); hl=QtWidgets.QVBoxLayout(self.home)
        title=QtWidgets.QLabel(f"<h1>{APP_TITLE} {APP_VERSION}</h1>")
        subtitle=QtWidgets.QLabel("<p>Auto mode runs: Hilbert (all categorical pairs, one-vs-rest), QDT batches, CbD from A1/A2/B1/B2, Quantum Game/Finance/Money per-row.</p>"); subtitle.setWordWrap(True)
        brand=QtWidgets.QLabel(f"<p><em>{BRAND}</em></p>"); brand.setStyleSheet("color:#555;")
        hl.addWidget(title); hl.addWidget(subtitle); hl.addWidget(brand); hl.addStretch(1)
        self.dataPanel=DataPanel(); self.specEditor=SpecEditor(); self.resultsPanel=ResultsPanel()
        self.autoPanel=AutoPanel(self.get_dataset); self.autoPanel.autoFinished.connect(self.on_result)
        self.runPanel=QtWidgets.QWidget(); rl=QtWidgets.QFormLayout(self.runPanel); self.runBtn=QtWidgets.QPushButton("Run Now"); self.runBtn.clicked.connect(self.run_model); rl.addRow(self.runBtn)
        for w in [self.home,self.dataPanel,self.specEditor,self.autoPanel,self.runPanel,self.resultsPanel]: self.stack.addWidget(w)
        self.stack.setCurrentWidget(self.home)
        self.btnHome.clicked.connect(lambda: self.stack.setCurrentWidget(self.home))
        self.btnData.clicked.connect(lambda: self.stack.setCurrentWidget(self.dataPanel))
        self.btnSpec.clicked.connect(lambda: self.stack.setCurrentWidget(self.specEditor))
        self.btnAuto.clicked.connect(lambda: self.stack.setCurrentWidget(self.autoPanel))
        self.btnRun.clicked.connect(lambda: self.stack.setCurrentWidget(self.runPanel))
        self.btnResults.clicked.connect(lambda: self.stack.setCurrentWidget(self.resultsPanel))
        self.specEditor.specChanged.connect(self.on_spec_changed)
        self.make_menu(); self.statusBar().showMessage(BRAND)
    def make_menu(self):
        m=self.menuBar(); fileM=m.addMenu("File")
        actNew=fileM.addAction("New Project"); actOpen=fileM.addAction("Open Project…"); actSave=fileM.addAction("Save Project…")
        fileM.addSeparator(); actQuit=fileM.addAction("Quit")
        actNew.triggered.connect(self.new_project); actOpen.triggered.connect(self.open_project); actSave.triggered.connect(self.save_project); actQuit.triggered.connect(self.close)
        helpM=m.addMenu("Help"); actAbout=helpM.addAction("About QuanEcon Studio"); actAbout.triggered.connect(self.about)
        # Extras
        samplesM=fileM.addMenu('Open Sample Dataset')
        actHS=samplesM.addAction('Hilbert / QDT / CbD / Finance / Money / Game…')
        actHS.triggered.connect(lambda: open_samples_menu(self, self._load_sample))
        helpM.addSeparator()
        actLearn=helpM.addAction('Learning Center (Six Methods)…')
        actLearn.triggered.connect(lambda: open_learning_dialog(self))
        actDocs=helpM.addAction('Built-in Docs (EULA, Manual, Privacy)…')
        actDocs.triggered.connect(lambda: open_docs_dialog(self))
    def _load_sample(self, path: str):
        try:
            self.dataPanel.load_path(path)
            self.stack.setCurrentWidget(self.dataPanel)
            self.statusBar().showMessage(f'Sample loaded: {path}', 5000)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, 'Sample Error', str(e))

    def about(self):
        QtWidgets.QMessageBox.information(
            self,
            'About',
            f"{APP_TITLE} {APP_VERSION}\n{BRAND}\n\nEmail: niazmirza111@gmail.com\nThis build includes: DOCX export, Learning Center, built-in Docs (EULA/Manual/Privacy),\nand one-click example datasets."
        )

    def get_dataset(self): return self.dataPanel.dataset
    def on_spec_changed(self,spec:Spec): self.spec=spec
    def run_model(self):
        if not self.spec: QtWidgets.QMessageBox.information(self,"No spec","Create a model spec first."); return
        self.runBtn.setEnabled(False); QtWidgets.QApplication.setOverrideCursor(QtCore.Qt.WaitCursor)
        try:
            runner=Runner(self.spec); runner.finished.connect(self.on_result); runner.run()
        finally:
            QtWidgets.QApplication.restoreOverrideCursor(); self.runBtn.setEnabled(True)
    def on_result(self,result:RunResult):
        self.result=result; self.resultsPanel.show_result(self.spec,result); self.stack.setCurrentWidget(self.resultsPanel)
        if not result.ok: QtWidgets.QMessageBox.warning(self,"Run info", result.summary)
    def new_project(self):
        self.dataset=None; self.spec=Spec(framework="Hilbert", config={}); self.result=None; self.specEditor.template("Hilbert"); self.stack.setCurrentWidget(self.home)
    def save_project(self):
        path,_=QtWidgets.QFileDialog.getSaveFileName(self,"Save Project","quanecon_project.json","JSON (*.json)")
        if not path: return
        proj={"dataset": self.dataset.name if self.dataset else None, "spec":{"framework":self.spec.framework,"config":self.spec.config}, "result": self.result.details if self.result else None}
        with open(path,'w',encoding='utf-8') as f: json.dump(proj,f,indent=2)
        QtWidgets.QMessageBox.information(self,"Saved", f"Project saved to:\n{path}")
    def open_project(self):
        path,_=QtWidgets.QFileDialog.getOpenFileName(self,"Open Project","","JSON (*.json)")
        if not path: return
        try:
            with open(path,'r',encoding='utf-8') as f: proj=json.load(f)
            spec=proj.get("spec",{}); fw=spec.get("framework","Hilbert"); cfg=spec.get("config",{})
            self.specEditor.box.setCurrentText(fw); self.specEditor.text.setPlainText(json.dumps(cfg, indent=2)); self.specEditor.emit()
            QtWidgets.QMessageBox.information(self,"Loaded", f"Project loaded from:\n{path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self,"Error", str(e))

def main():
    app=QtWidgets.QApplication(sys.argv); w=MainWindow(); w.show(); sys.exit(app.exec())

if __name__=="__main__":
    main()
